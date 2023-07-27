def combine(para1,para2,name):
# Create a new Word document
    from docx import Document
    import os
    doc = Document()
    doc.add_paragraph("Transcript : ")

            # Add the first paragraph
    doc.add_paragraph(para1)

            # Add two line spaces
    doc.add_paragraph()
    doc.add_paragraph("Summary : ")

            # Add the second paragraph
    doc.add_paragraph(para2)

    file_path = os.path.join("C:/Users/sumit/Summarizer/media/new_transcripts/", name)
    doc.save(file_path)
